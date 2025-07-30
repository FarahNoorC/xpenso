import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import io
from docx import Document
from docx.shared import Inches
from PIL import Image

st.title("📊 Flexible Chart Generator")

uploaded_file = st.file_uploader("Upload CSV or Excel file", type=["csv", "xlsx"])

if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1]

    # Read the uploaded file
    if file_type == "csv":
        df = pd.read_csv(uploaded_file)
    elif file_type == "xlsx":
        xls = pd.ExcelFile(uploaded_file)
        sheet_name = st.selectbox("Select Excel sheet", xls.sheet_names)
        df = xls.parse(sheet_name)

    st.subheader("🔍 Data Preview")
    st.dataframe(df)

    x_col = st.selectbox("Select X-axis column", df.columns)
    y_col = st.selectbox("Select Y-axis column", df.columns)

    row_start = st.number_input("Start row (0-based)", min_value=0, max_value=len(df)-1, value=0)
    row_end = st.number_input("End row (exclusive)", min_value=row_start+1, max_value=len(df), value=len(df))

    chart_type = st.selectbox("Select chart type", ["line", "bar", "scatter", "pie", "histogram"])
    df_subset = df.iloc[int(row_start):int(row_end)]

    st.subheader(f"📈 {chart_type.capitalize()} Chart")
    fig, ax = plt.subplots(figsize=(10, 6))

    if chart_type == "line":
        sns.lineplot(data=df_subset, x=x_col, y=y_col, ax=ax)
    elif chart_type == "bar":
        sns.barplot(data=df_subset, x=x_col, y=y_col, ax=ax)
    elif chart_type == "scatter":
        sns.scatterplot(data=df_subset, x=x_col, y=y_col, ax=ax)
    elif chart_type == "pie":
        df_subset = df_subset.set_index(x_col)
        df_subset[y_col].plot.pie(autopct="%1.1f%%", ax=ax)
        ax.set_ylabel("")
    elif chart_type == "histogram":
        sns.histplot(data=df_subset, x=y_col, kde=True, ax=ax)

    plt.xticks(rotation=45)
    st.pyplot(fig)

    # --------------------- PDF Download ---------------------
    pdf_bytes = io.BytesIO()
    fig.savefig(pdf_bytes, format='pdf')
    pdf_bytes.seek(0)

    st.download_button(
        label="📄 Download Chart as PDF",
        data=pdf_bytes,
        file_name="chart.pdf",
        mime="application/pdf"
    )

    # --------------------- Word Download ---------------------
    img_bytes = io.BytesIO()
    fig.savefig(img_bytes, format='png')
    img_bytes.seek(0)

    doc = Document()
    doc.add_heading(f'{chart_type.capitalize()} Chart', level=1)
    doc.add_picture(img_bytes, width=Inches(6))

    word_bytes = io.BytesIO()
    doc.save(word_bytes)
    word_bytes.seek(0)

    st.download_button(
        label="📝 Download Chart as Word Document",
        data=word_bytes,
        file_name="chart.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
