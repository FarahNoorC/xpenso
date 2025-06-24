import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO


def generate_word_report(df, sheet_name="Sheet1"):
    doc = Document()
    doc.add_heading(f'Report: {sheet_name}', 0)
    doc.add_paragraph(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")

    # Column overview
    doc.add_heading('Column Overview', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Column'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Missing Values'

    for col in df.columns:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = str(df[col].dtype)
        row_cells[2].text = str(df[col].isnull().sum())

    # Numeric stats
    num_cols = df.select_dtypes(include=['number']).columns
    numeric_df = df[num_cols].apply(pd.to_numeric, errors='coerce')
    
    if not numeric_df.empty:
        doc.add_heading('Descriptive Statistics (Numeric)', level=1)
        stats = numeric_df.describe().round(2)
        totals = numeric_df.sum().round(2)
        stats.loc['total'] = totals
        doc.add_paragraph(stats.to_string())

    # Categorical stats
    cat_cols = df.select_dtypes(include=['object', 'category']).columns
    if not cat_cols.empty:
        doc.add_heading('Top Unique Values (Categorical)', level=1)
        for col in cat_cols:
            doc.add_paragraph(f"{col}:", style='List Bullet')
            vc = df[col].astype(str).value_counts().head(5)
            for val, count in vc.items():
                doc.add_paragraph(f"{val}: {count}", style='List Number')

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Streamlit App ---
st.title("📄 Excel/CSV → Word Report Generator with Totals")

uploaded_file = st.file_uploader("Upload Excel or CSV file", type=["csv", "xlsx"])

if uploaded_file:
    file_ext = uploaded_file.name.split(".")[-1].lower()

    if file_ext == "csv":
        df = pd.read_csv(uploaded_file)
        sheet_name = "CSV"
    else:
        xls = pd.ExcelFile(uploaded_file)
        sheet_name = st.selectbox("Select Sheet", xls.sheet_names)
        df = xls.parse(sheet_name)

    st.subheader("Data Preview")
    st.dataframe(df.head())

    if st.button("Generate Word Report"):
        word_file = generate_word_report(df, sheet_name)
        st.download_button("📥 Download Word Report", word_file, file_name="report.docx")
